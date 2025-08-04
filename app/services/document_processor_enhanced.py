#!/usr/bin/env python3
"""
Enhanced Document Processor baseado nos padrões AGNO Framework
Implementa processamento multimodal avançado com OCR e múltiplos formatos
"""

import io
import base64
from typing import Dict, Any, Optional, Union, List
from pathlib import Path
from PIL import Image as PILImage
import pytesseract
from pypdf import PdfReader
import docx
from loguru import logger

from app.utils.logger import emoji_logger


class EnhancedDocumentProcessor:
    """
    Processador de documentos aprimorado baseado nos padrões AGNO Framework
    
    Suporta:
    - PDFs com OCR de imagens embarcadas
    - Documentos Word (.docx)
    - Processamento direto de URLs
    - Múltiplos formatos de input (path, bytes, file-like)
    """
    
    def __init__(self):
        """Initialize o processador"""
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_doc_fallback,
            'txt': self._process_txt,
            'rtf': self._process_rtf_fallback
        }
        
        # Magic bytes expandidos baseados no AGNO
        self.magic_bytes = {
            b'%PDF': 'pdf',
            b'PK\x03\x04': 'docx',  # ZIP-based (docx, xlsx, etc)
            b'\xd0\xcf\x11\xe0': 'doc',  # Old Office format
            b'{\rtf': 'rtf',
            b'{\\\rtf': 'rtf'
        }
    
    async def process_document(
        self, 
        data: Union[str, bytes, io.BytesIO], 
        filename: Optional[str] = None,
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Processa documento com detecção automática de formato
        
        Args:
            data: Dados do documento (base64, bytes, ou file-like)
            filename: Nome do arquivo (para detecção de formato)
            enable_ocr: Habilitar OCR para imagens em PDFs
            
        Returns:
            Dict com conteúdo extraído e metadados
        """
        try:
            emoji_logger.system_info("Iniciando processamento de documento avançado")
            
            # Converter base64 para bytes se necessário
            if isinstance(data, str):
                try:
                    document_bytes = base64.b64decode(data)
                    emoji_logger.system_info(f"Base64 decodificado: {len(document_bytes)} bytes")
                except Exception as decode_error:
                    emoji_logger.system_error("Base64 Decode", str(decode_error))
                    return {
                        "type": "document",
                        "error": "Erro ao decodificar base64",
                        "status": "error"
                    }
            elif isinstance(data, bytes):
                document_bytes = data
            else:
                # Assume file-like object
                document_bytes = data.read()
            
            # Detectar formato do documento
            document_format = self._detect_format(document_bytes, filename)
            emoji_logger.system_info(f"Formato detectado: {document_format}")
            
            if document_format not in self.supported_formats:
                emoji_logger.system_warning(f"Formato não suportado: {document_format}")
                return {
                    "type": "document",
                    "filename": filename or "documento",
                    "error": f"Formato não suportado: {document_format}",
                    "status": "unsupported",
                    "supported_formats": list(self.supported_formats.keys())
                }
            
            # Processar documento com método específico
            processor = self.supported_formats[document_format]
            result = await processor(document_bytes, filename, enable_ocr)
            
            # Adicionar metadados padrão AGNO
            result.update({
                "processor": "enhanced_document_processor",
                "format": document_format,
                "ocr_enabled": enable_ocr,
                "processing_success": True
            })
            
            emoji_logger.system_info(f"Documento processado: {len(result.get('content', ''))} caracteres extraídos")
            return result
            
        except Exception as e:
            emoji_logger.system_error("Document Processing", str(e))
            return {
                "type": "document",
                "filename": filename or "documento",
                "error": f"Erro no processamento: {str(e)}",
                "status": "error",
                "processing_success": False
            }
    
    def _detect_format(self, data: bytes, filename: Optional[str] = None) -> str:
        """
        Detecta formato do documento usando magic bytes e extensão
        Baseado nos padrões AGNO Framework
        """
        # Tentar magic bytes primeiro (mais confiável)
        for magic, format_type in self.magic_bytes.items():
            if data.startswith(magic):
                return format_type
        
        # Fallback para extensão do arquivo
        if filename:
            extension = Path(filename).suffix.lower().lstrip('.')
            if extension in self.supported_formats:
                return extension
        
        # Default para texto se não conseguir detectar
        return 'txt'
    
    async def _process_pdf(
        self, 
        data: bytes, 
        filename: Optional[str] = None, 
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Processa PDF com OCR opcional (padrão AGNO PDFImageReader)
        """
        try:
            pdf_stream = io.BytesIO(data)
            reader = PdfReader(pdf_stream)
            
            extracted_text = ""
            images_processed = 0
            ocr_text = ""
            
            # Extrair texto de todas as páginas
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        extracted_text += f"\n--- Página {page_num + 1} ---\n{page_text}\n"
                    
                    # OCR de imagens se habilitado (padrão AGNO)
                    if enable_ocr and '/XObject' in page.get('/Resources', {}):
                        xObject = page['/Resources']['/XObject'].get_object()
                        
                        for obj in xObject:
                            if xObject[obj]['/Subtype'] == '/Image':
                                try:
                                    # Extrair dados da imagem
                                    size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
                                    data_stream = xObject[obj].get_data()
                                    
                                    # Tentar OCR na imagem
                                    if data_stream:
                                        img = PILImage.open(io.BytesIO(data_stream))
                                        img_text = pytesseract.image_to_string(img, lang='por')
                                        
                                        if img_text.strip():
                                            ocr_text += f"\n[OCR Imagem Página {page_num + 1}]: {img_text}\n"
                                            images_processed += 1
                                            
                                except Exception as ocr_error:
                                    logger.debug(f"OCR falhou na página {page_num + 1}: {ocr_error}")
                                    continue
                                    
                except Exception as page_error:
                    logger.debug(f"Erro na página {page_num + 1}: {page_error}")
                    continue
            
            # Combinar texto extraído e OCR
            full_content = extracted_text
            if ocr_text.strip():
                full_content += f"\n--- Conteúdo de Imagens (OCR) ---\n{ocr_text}"
            
            result = {
                "type": "document",
                "filename": filename or "documento.pdf",
                "content": full_content.strip(),
                "pages": len(reader.pages),
                "text_extracted": bool(extracted_text.strip()),
                "images_processed": images_processed,
                "ocr_content": bool(ocr_text.strip()),
                "status": "success"
            }
            
            emoji_logger.system_info(
                f"PDF processado: {len(reader.pages)} páginas, "
                f"{images_processed} imagens com OCR"
            )
            
            return result
            
        except Exception as e:
            emoji_logger.system_error("PDF Processing", str(e))
            return {
                "type": "document",
                "filename": filename or "documento.pdf",
                "error": f"Erro ao processar PDF: {str(e)}",
                "status": "error"
            }
    
    async def _process_docx(
        self, 
        data: bytes, 
        filename: Optional[str] = None, 
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Processa documento Word (.docx) - padrão AGNO DocxReader
        """
        try:
            docx_stream = io.BytesIO(data)
            doc = docx.Document(docx_stream)
            
            # Extrair texto dos parágrafos
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extrair texto das tabelas
            tables_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_content.append(row_text)
            
            # Combinar conteúdo
            full_content = "\n\n".join(paragraphs)
            if tables_content:
                full_content += "\n\n--- Tabelas ---\n" + "\n".join(tables_content)
            
            return {
                "type": "document",
                "filename": filename or "documento.docx",
                "content": full_content.strip(),
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "status": "success"
            }
            
        except Exception as e:
            emoji_logger.system_error("DOCX Processing", str(e))
            return {
                "type": "document",
                "filename": filename or "documento.docx",
                "error": f"Erro ao processar DOCX: {str(e)}",
                "status": "error"
            }
    
    async def _process_doc_fallback(
        self, 
        data: bytes, 
        filename: Optional[str] = None, 
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Fallback para documentos .doc antigos
        """
        try:
            # Para .doc antigos, seria necessário uma lib como python-docx2txt
            # Por enquanto, retornamos indicação de formato não suportado completamente
            return {
                "type": "document",
                "filename": filename or "documento.doc",
                "content": "[Documento .doc detectado - conversão limitada disponível]",
                "error": "Formato .doc requer conversão prévia para .docx",
                "status": "partial_support"
            }
            
        except Exception as e:
            emoji_logger.system_error("DOC Processing", str(e))
            return {
                "type": "document",
                "filename": filename or "documento.doc",
                "error": f"Erro ao processar DOC: {str(e)}",
                "status": "error"
            }
    
    async def _process_txt(
        self, 
        data: bytes, 
        filename: Optional[str] = None, 
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Processa arquivo de texto simples
        """
        try:
            # Tentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            content = None
            for encoding in encodings:
                try:
                    content = data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Não foi possível decodificar o arquivo de texto")
            
            return {
                "type": "document",
                "filename": filename or "documento.txt",
                "content": content.strip(),
                "encoding": encoding,
                "status": "success"
            }
            
        except Exception as e:
            emoji_logger.system_error("TXT Processing", str(e))
            return {
                "type": "document",
                "filename": filename or "documento.txt",
                "error": f"Erro ao processar TXT: {str(e)}",
                "status": "error"
            }
    
    async def _process_rtf_fallback(
        self, 
        data: bytes, 
        filename: Optional[str] = None, 
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Fallback básico para documentos RTF
        """
        try:
            # RTF básico - extração simples
            content = data.decode('utf-8', errors='ignore')
            
            # Remover códigos RTF básicos (muito simplificado)
            import re
            clean_text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            clean_text = re.sub(r'[{}\\]', '', clean_text)
            
            return {
                "type": "document",
                "filename": filename or "documento.rtf",
                "content": clean_text.strip(),
                "status": "basic_extraction"
            }
            
        except Exception as e:
            emoji_logger.system_error("RTF Processing", str(e))
            return {
                "type": "document",
                "filename": filename or "documento.rtf",
                "error": f"Erro ao processar RTF: {str(e)}",
                "status": "error"
            }


# Instância global do processador aprimorado
enhanced_processor = EnhancedDocumentProcessor()


async def process_document_enhanced(
    data: Union[str, bytes, io.BytesIO],
    filename: Optional[str] = None,
    enable_ocr: bool = True
) -> Dict[str, Any]:
    """
    Interface pública para processamento de documentos
    
    Args:
        data: Dados do documento
        filename: Nome do arquivo
        enable_ocr: Habilitar OCR para PDFs com imagens
        
    Returns:
        Resultado do processamento
    """
    return await enhanced_processor.process_document(data, filename, enable_ocr)