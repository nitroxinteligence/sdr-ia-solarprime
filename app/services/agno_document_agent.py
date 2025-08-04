#!/usr/bin/env python3
"""
AGNO DocumentEnhancementAgent - Sub-agente modular para processamento de documentos
Baseado nos padrões AGNO Framework: PDFReader, PDFImageReader, DocxReader

FILOSOFIA: ZERO COMPLEXIDADE - O SIMPLES FUNCIONA
- Wrappeia o código existente sem breaking changes
- Adiciona OCR para PDFs com imagens (AGNO PDFImageReader pattern)
- Adiciona suporte DOCX (AGNO DocxReader pattern)
- Decorator pattern para modularidade
"""

import base64
import io
from typing import Dict, Any, Optional, Callable, List
from pathlib import Path
from PIL import Image as PILImage
from app.utils.logger import emoji_logger

# Fallback imports caso algumas libs não estejam disponíveis
try:
    import fitz  # PyMuPDF para OCR de imagens em PDFs
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    emoji_logger.system_warning("PyMuPDF não disponível - usando pypdf como fallback")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    emoji_logger.system_warning("python-docx não disponível - suporte DOCX limitado")

try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    emoji_logger.system_warning("pytesseract não disponível - OCR desabilitado")


class AGNODocumentProcessor:
    """
    Processador de documentos baseado nos padrões AGNO Framework
    - PDFReader: Extração básica de texto
    - PDFImageReader: PDF + OCR de imagens embarcadas
    - DocxReader: Processamento de documentos Word
    """
    
    def __init__(self):
        """Initialize com padrões AGNO"""
        self.supported_formats = {
            'pdf': self._enhance_pdf_processing,
            'docx': self._process_docx_agno,
            'doc': self._process_doc_fallback,
            'txt': self._process_txt_agno
        }
        
        # Magic bytes para detecção de formato (padrão AGNO)
        self.document_magic_bytes = {
            b'%PDF': 'pdf',
            b'PK\x03\x04': 'docx',  # ZIP-based (docx, xlsx, etc)
            b'\xd0\xcf\x11\xe0': 'doc',  # Old Office format
        }
    
    def enhance_document_processing(
        self, 
        document_bytes: bytes, 
        filename: Optional[str] = None,
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Enhanced document processing com AGNO patterns
        
        Args:
            document_bytes: Bytes do documento
            filename: Nome do arquivo
            enable_ocr: Habilitar OCR (AGNO PDFImageReader)
            
        Returns:
            Dict com conteúdo extraído + AGNO metadata
        """
        try:
            # Detectar formato do documento
            detected_format = self._detect_document_format(document_bytes, filename)
            emoji_logger.system_info(f"AGNO Document format detected: {detected_format}")
            
            if detected_format not in self.supported_formats:
                return {
                    'type': 'document',
                    'error': f'Formato não suportado pelo AGNO: {detected_format}',
                    'agno_compatible': False,
                    'supported_formats': list(self.supported_formats.keys())
                }
            
            # Processar com método específico AGNO
            processor = self.supported_formats[detected_format]
            result = processor(document_bytes, filename, enable_ocr)
            
            # Adicionar metadata AGNO padrão
            result.update({
                'agno_processor': 'document_enhancement_agent',
                'detected_format': detected_format,
                'ocr_enabled': enable_ocr and detected_format == 'pdf',
                'agno_compatible': True,
                'processing_success': True
            })
            
            return result
            
        except Exception as e:
            emoji_logger.system_error("AGNO Document Enhancement", str(e))
            return {
                'type': 'document',
                'error': f'Erro no processamento AGNO: {str(e)}',
                'agno_compatible': False
            }
    
    def _detect_document_format(self, data: bytes, filename: Optional[str] = None) -> str:
        """Detecta formato usando magic bytes + filename (padrão AGNO)"""
        # Magic bytes primeiro (mais confiável)
        for magic, format_type in self.document_magic_bytes.items():
            if data.startswith(magic):
                return format_type
        
        # Fallback para extensão
        if filename:
            extension = Path(filename).suffix.lower().lstrip('.')
            if extension in self.supported_formats:
                return extension
        
        return 'txt'  # Default
    
    def _enhance_pdf_processing(
        self, 
        document_bytes: bytes, 
        filename: Optional[str] = None,
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Enhanced PDF processing - AGNO PDFImageReader pattern
        Adiciona OCR de imagens embarcadas ao processamento básico
        """
        try:
            pdf_stream = io.BytesIO(document_bytes)
            
            # Usar PyMuPDF se disponível, senão fallback para pypdf
            if PYMUPDF_AVAILABLE:
                return self._process_pdf_with_pymupdf(pdf_stream, filename, enable_ocr)
            else:
                return self._process_pdf_with_pypdf(pdf_stream, filename, enable_ocr)
                
        except Exception as e:
            emoji_logger.system_error("AGNO PDF Processing", str(e))
            return {
                'type': 'document',
                'filename': filename or 'documento.pdf',
                'error': f'Erro no processamento PDF AGNO: {str(e)}',
                'status': 'error'
            }
    
    def _process_pdf_with_pymupdf(self, pdf_stream: io.BytesIO, filename: Optional[str], enable_ocr: bool) -> Dict[str, Any]:
        """PyMuPDF processing with OCR"""
        doc = fitz.open(stream=pdf_stream, filetype="pdf")
        
        extracted_text = ""
        images_processed = 0
        ocr_text = ""
        total_pages = len(doc)
        
        for page_num in range(total_pages):
            page = doc[page_num]
            
            # Extrair texto básico
            page_text = page.get_text()
            if page_text.strip():
                extracted_text += f"\n--- Página {page_num + 1} ---\n{page_text}\n"
            
            # OCR de imagens se habilitado (AGNO PDFImageReader pattern)
            if enable_ocr and OCR_AVAILABLE:
                try:
                    # Extrair imagens da página
                    image_list = page.get_images()
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            # Extrair dados da imagem
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            
                            if pix.n - pix.alpha < 4:  # GRAY ou RGB
                                img_data = pix.tobytes("png")
                                
                                # OCR na imagem
                                pil_img = PILImage.open(io.BytesIO(img_data))
                                img_text = pytesseract.image_to_string(pil_img, lang='por')
                                
                                if img_text.strip():
                                    ocr_text += f"\n[OCR Imagem {img_index + 1} - Página {page_num + 1}]: {img_text.strip()}\n"
                                    images_processed += 1
                            
                            pix = None  # Liberar memória
                            
                        except Exception as img_error:
                            emoji_logger.system_warning(f"OCR falhou para imagem {img_index}: {img_error}")
                            continue
                            
                except Exception as ocr_error:
                    emoji_logger.system_warning(f"OCR falhou na página {page_num + 1}: {ocr_error}")
                    continue
        
        doc.close()
        
        # Combinar texto extraído + OCR
        full_content = extracted_text
        if ocr_text.strip():
            full_content += f"\n--- Conteúdo de Imagens (OCR AGNO) ---\n{ocr_text}"
        
        result = {
            'type': 'document',
            'filename': filename or 'documento.pdf',
            'content': full_content.strip(),
            'pages': total_pages,
            'text_extracted': bool(extracted_text.strip()),
            'images_processed': images_processed,
            'ocr_content': bool(ocr_text.strip()),
            'agno_pattern': 'PDFImageReader' if enable_ocr else 'PDFReader',
            'status': 'success'
        }
        
        emoji_logger.system_info(
            f"AGNO PDF enhanced: {total_pages} páginas, {images_processed} imagens com OCR"
        )
        
        return result
    
    def _process_pdf_with_pypdf(self, pdf_stream: io.BytesIO, filename: Optional[str], enable_ocr: bool) -> Dict[str, Any]:
        """Fallback PDF processing usando pypdf (sem OCR de imagens)"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(pdf_stream)
            extracted_text = ""
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        extracted_text += f"\n--- Página {page_num + 1} ---\n{page_text}\n"
                except Exception as page_error:
                    emoji_logger.system_warning(f"Erro na página {page_num + 1}: {page_error}")
                    continue
            
            result = {
                'type': 'document',
                'filename': filename or 'documento.pdf',
                'content': extracted_text.strip(),
                'pages': len(reader.pages),
                'text_extracted': bool(extracted_text.strip()),
                'images_processed': 0,  # pypdf não faz OCR
                'ocr_content': False,
                'agno_pattern': 'PDFReader_fallback',
                'status': 'success',
                'fallback_used': 'pypdf'
            }
            
            emoji_logger.system_info(f"PDF processado com pypdf fallback: {len(reader.pages)} páginas")
            
            return result
            
        except Exception as e:
            emoji_logger.system_error("AGNO PDF Fallback", str(e))
            return {
                'type': 'document',
                'filename': filename or 'documento.pdf',
                'error': f'Erro no fallback PDF: {str(e)}',
                'status': 'error'
            }
    
    def _process_docx_agno(
        self, 
        document_bytes: bytes, 
        filename: Optional[str] = None,
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        DOCX processing - AGNO DocxReader pattern
        """
        if not DOCX_AVAILABLE:
            return {
                'type': 'document',
                'filename': filename or 'documento.docx',
                'error': 'python-docx não disponível - instale com: pip install python-docx',
                'status': 'library_missing'
            }
        
        try:
            docx_stream = io.BytesIO(document_bytes)
            doc = docx.Document(docx_stream)
            
            # Extrair parágrafos
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extrair tabelas
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_data.append(row_text)
                if table_data:
                    tables_content.extend(table_data)
            
            # Combinar conteúdo
            full_content = "\n\n".join(paragraphs)
            if tables_content:
                full_content += "\n\n--- Tabelas ---\n" + "\n".join(tables_content)
            
            result = {
                'type': 'document',
                'filename': filename or 'documento.docx',
                'content': full_content.strip(),
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
                'agno_pattern': 'DocxReader',
                'status': 'success'
            }
            
            emoji_logger.system_info(f"AGNO DOCX processed: {len(paragraphs)} parágrafos, {len(doc.tables)} tabelas")
            return result
            
        except Exception as e:
            emoji_logger.system_error("AGNO DOCX Processing", str(e))
            return {
                'type': 'document',
                'filename': filename or 'documento.docx',
                'error': f'Erro no processamento DOCX AGNO: {str(e)}',
                'status': 'error'
            }
    
    def _process_doc_fallback(
        self, 
        document_bytes: bytes, 
        filename: Optional[str] = None,
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """Fallback para DOC antigos"""
        return {
            'type': 'document',
            'filename': filename or 'documento.doc',
            'content': '[Documento .doc detectado - requer conversão para .docx para processamento completo]',
            'error': 'Formato .doc legado - converta para .docx para melhor suporte',
            'agno_pattern': 'LegacyDocFallback',
            'status': 'partial_support'
        }
    
    def _process_txt_agno(
        self, 
        document_bytes: bytes, 
        filename: Optional[str] = None,
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """TXT processing com encoding detection"""
        try:
            # Tentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    content = document_bytes.decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Não foi possível decodificar o arquivo de texto")
            
            return {
                'type': 'document',
                'filename': filename or 'documento.txt',
                'content': content.strip(),
                'encoding': used_encoding,
                'agno_pattern': 'TextReader',
                'status': 'success'
            }
            
        except Exception as e:
            emoji_logger.system_error("AGNO TXT Processing", str(e))
            return {
                'type': 'document',
                'filename': filename or 'documento.txt',
                'error': f'Erro no processamento TXT AGNO: {str(e)}',
                'status': 'error'
            }


# Instância global do processador AGNO
agno_document_processor = AGNODocumentProcessor()


def agno_document_enhancer(func: Callable) -> Callable:
    """
    Decorator AGNO para enhanced document processing
    Wrappeia método existente + adiciona AGNO capabilities (OCR, DOCX, etc)
    
    ZERO COMPLEXIDADE: Mantém código existente funcionando
    """
    def wrapper(*args, **kwargs):
        # Chama método original (mantém funcionalidade existente)
        original_result = func(*args, **kwargs)
        
        # Se processamento original foi bem-sucedido E é documento, aplicar AGNO enhancements
        if (isinstance(original_result, dict) and 
            original_result.get('type') == 'document' and 
            not original_result.get('error')):
            
            try:
                # Extrair document_bytes do contexto
                if len(args) >= 2 and isinstance(args[1], str):
                    # args[1] provavelmente é media_data (base64)
                    document_bytes = base64.b64decode(args[1])
                    filename = kwargs.get('filename') or original_result.get('filename')
                    
                    # Aplicar AGNO enhancements
                    agno_result = agno_document_processor.enhance_document_processing(
                        document_bytes, 
                        filename, 
                        enable_ocr=True
                    )
                    
                    # Se AGNO conseguiu processar melhor, usar o resultado AGNO
                    if (agno_result.get('processing_success') and 
                        len(agno_result.get('content', '')) > len(original_result.get('content', ''))):
                        
                        # Combinar resultado original + AGNO enhancements
                        enhanced_result = original_result.copy()
                        enhanced_result.update({
                            'content': agno_result['content'],
                            'agno_metadata': agno_result,
                            'enhanced_with_agno': True,
                            'processing_method': 'agno_enhanced'
                        })
                        
                        emoji_logger.system_info(
                            f"AGNO Document enhanced: {agno_result.get('agno_pattern', 'unknown')} "
                            f"- {len(agno_result.get('content', ''))} caracteres"
                        )
                        
                        return enhanced_result
                    else:
                        # AGNO não melhorou significativamente, manter original + metadata
                        original_result['agno_metadata'] = agno_result
                        original_result['agno_enhancement_applied'] = False
                        
            except Exception as e:
                emoji_logger.system_warning(f"AGNO document enhancement failed (non-critical): {e}")
                # Não falha - apenas não adiciona enhancements
                original_result['agno_enhancement_error'] = str(e)
        
        # Se é documento mas não foi processado (ex: formato não suportado), tentar AGNO fallback
        elif (isinstance(original_result, dict) and 
              original_result.get('type') == 'document' and 
              original_result.get('error')):
            
            try:
                if len(args) >= 2 and isinstance(args[1], str):
                    document_bytes = base64.b64decode(args[1])
                    filename = kwargs.get('filename') or original_result.get('filename')
                    
                    # Tentar AGNO fallback
                    agno_fallback = agno_document_processor.enhance_document_processing(
                        document_bytes, 
                        filename, 
                        enable_ocr=True
                    )
                    
                    # Se AGNO conseguiu processar onde o original falhou
                    if agno_fallback.get('processing_success'):
                        agno_fallback['fallback_from_original'] = True
                        agno_fallback['original_error'] = original_result.get('error')
                        
                        emoji_logger.system_info(f"AGNO fallback successful for: {filename}")
                        return agno_fallback
                        
            except Exception as e:
                emoji_logger.system_warning(f"AGNO fallback failed: {e}")
        
        return original_result
    
    return wrapper